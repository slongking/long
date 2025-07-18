"""
社交媒体舆情分析系统 v2.1
功能特性：
1. 全流程异常处理机制
2. 自适应文本清洗
3. 多维度可视化分析
4. 自动生成交互式报告
5. 数据质量监控体系
"""

import pandas as pd
import sqlalchemy
import matplotlib.pyplot as plt
import seaborn as sns
from wordcloud import WordCloud
from snownlp import SnowNLP
import plotly.express as px
from plotly.subplots import make_subplots
import plotly.graph_objects as go
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
import jieba
import logging
import platform
from pathlib import Path
from typing import Dict, Any, Optional
from concurrent.futures import ThreadPoolExecutor, as_completed
from tqdm import tqdm
import sys
import json
import re
from functools import lru_cache
import numpy as np

# 日志配置
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('analysis.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger('PublicOpinionAnalysis')

class AnalysisConfig:
    """系统配置中心"""
    def __init__(self):
        self.output_dir = Path('output')
        self.cache_dir = Path('cache')
        self.font_path = self._get_font_path()
        self.sentiment_bins = [0, 0.3, 0.7, 1]
        self._init_directories()
        
    def _get_font_path(self) -> str:
        """自动获取系统字体路径"""
        system = platform.system()
        try:
            if system == 'Windows':
                return 'C:/Windows/Fonts/simhei.ttf'
            elif system == 'Linux':
                return '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc'
            else:
                return 'Arial Unicode MS'
        except Exception as e:
            logger.warning(f"字体加载失败: {str(e)}，使用备用字体")
            return None
            
    def _init_directories(self):
        """创建必要目录"""
        self.output_dir.mkdir(exist_ok=True)
        self.cache_dir.mkdir(exist_ok=True)

class DataProcessor:
    """数据处理器"""
    def __init__(self, db_config: dict):
        self.db_config = db_config
        self.config = AnalysisConfig()
        self.df = None
        
    def _create_engine(self):
        """创建数据库引擎"""
        try:
            return sqlalchemy.create_engine(
                f"mysql+pymysql://{self.db_config['user']}:{self.db_config['password']}"
                f"@{self.db_config['host']}:{self.db_config['port']}/{self.db_config['db']}",
                pool_size=5,
                max_overflow=10,
                pool_pre_ping=True,
                connect_args={'connect_timeout': 10}
            )
        except Exception as e:
            logger.error(f"数据库引擎创建失败: {str(e)}")
            raise
            
    def fetch_data(self) -> pd.DataFrame:
        """安全获取数据"""
        try:
            engine = self._create_engine()
            with engine.connect() as conn:
                logger.info("成功建立数据库连接")
                return pd.read_sql(
                    "SELECT * FROM xhs_comment WHERE CommentTime > %(start_date)s",
                    conn,
                    params={'start_date': datetime.now() - timedelta(days=30)}
                )
        except Exception as e:
            logger.error(f"数据获取失败: {str(e)}")
            raise
            
    def _clean_text(self, text: str) -> str:
        """深度文本清洗"""
        if not isinstance(text, str):
            return ''
            
        # 移除特殊字符
        text = re.sub(r'[^\w\u4e00-\u9fff]', ' ', text)
        # 合并空白字符
        text = re.sub(r'\s+', ' ', text).strip()
        # 过滤无效文本
        if len(text) < 2 or text.isnumeric():
            return ''
        return text
            
    def preprocess(self) -> pd.DataFrame:
        """数据预处理流水线"""
        try:
            logger.info("启动数据预处理")
            raw_df = self.fetch_data()
            
            # 空值处理
            fill_values = {
                'CommentText': '',
                'CommentIp': '未知',
                'CommentLike': 0,
                'CommentUserId': '匿名用户'
            }
            df = raw_df.fillna(fill_values)
            
            # 类型转换
            time_cols = ['CommentTime', 'ColTime']
            df[time_cols] = df[time_cols].apply(pd.to_datetime, errors='coerce')
            df = df[~df['CommentText'].str.contains('R', na=False)]
            df = df[~df['CommentText'].str.contains('doge', na=False)]
            df = df[~df['CommentText'].str.contains('加入该群', na=False)]
            
            # 数据过滤
            df = df[
                (df['CommentTime'] >= pd.Timestamp.now() - pd.DateOffset(months=1)) &
                (df['CommentTime'] <= pd.Timestamp.now())
            ]
            
            # 文本清洗
            df['clean_text'] = df['CommentText'].apply(self._clean_text)
            
            # 数据质量检查
            if df.empty:
                raise ValueError("清洗后数据为空，请检查数据源")
                
            self.df = df
            logger.info(f"有效数据量: {len(df)} 条")
            return df
        except Exception as e:
            logger.error(f"预处理失败: {str(e)}")
            raise
            # 筛除评论中包含 'R' 的评论
            

class SentimentAnalyzer:
    """增强版情感分析器"""
    def __init__(self):
        self.config = AnalysisConfig()
        
    @lru_cache(maxsize=10000)
    def _cached_sentiment(self, text: str) -> float:
        """带缓存的情感分析"""
        try:
            if not text or len(text) < 2:
                return 0.5
                
            s = SnowNLP(text)
            if len(s.words) < 1:
                return 0.5
                
            return s.sentiments
        except ZeroDivisionError:
            logger.warning(f"ZeroDivisionError处理文本: {text[:50]}")
            return 0.5
        except Exception as e:
            logger.error(f"情感分析异常: {str(e)} - 文本: {text[:50]}")
            return 0.5
            
    def analyze(self, df: pd.DataFrame) -> pd.DataFrame:
        """执行情感分析"""
        try:
            logger.info("启动情感分析")
            texts = df['clean_text'].tolist()
            
            # 并行处理
            with ThreadPoolExecutor() as executor:
                futures = [executor.submit(self._cached_sentiment, text) for text in texts]
                results = []
                for future in tqdm(as_completed(futures), total=len(texts), desc="情感分析进度"):
                    results.append(future.result())
                    
            df['sentiment'] = results
            
            # 动态阈值调整
            try:
                q30 = np.nanquantile(df['sentiment'], 0.3)
                q70 = np.nanquantile(df['sentiment'], 0.7)
                bins = [0, q30, q70, 1]
            except Exception as e:
                logger.warning(f"分位数计算异常，使用默认阈值: {str(e)}")
                bins = self.config.sentiment_bins
                
            # 分类处理
            labels = ['负面', '中性', '正面']
            df['sentiment_label'] = pd.cut(
                df['sentiment'],
                bins=bins,
                labels=labels,
                include_lowest=True
            )
            
            # 处理异常分类
            df['sentiment_label'] = df['sentiment_label'].cat.add_categories('未知').fillna('未知')
            
            logger.info("情感分析完成")
            return df
        except Exception as e:
            logger.error(f"情感分析失败: {str(e)}")
            raise

class VisualizationEngine:
    """可视化引擎"""
    def __init__(self):
        self.config = AnalysisConfig()
        
    def generate_wordcloud(self, df: pd.DataFrame):
        """生成词云"""
        try:
            logger.info("生成词云")
            text = ' '.join(df['clean_text'])
            if not text:
                logger.warning("无有效文本生成词云")
                return
                
            wc = WordCloud(
                font_path=self.config.font_path,
                width=1600,
                height=1200,
                background_color='white',
                collocations=False
            ).generate(text)
            
            plt.figure(figsize=(20, 15))
            plt.imshow(wc)
            plt.axis('off')
            plt.savefig(self.config.output_dir / 'wordcloud.png', bbox_inches='tight')
            plt.close()
        except Exception as e:
            logger.error(f"词云生成失败: {str(e)}")
            
    def create_dashboard(self, df: pd.DataFrame):
        """创建交互式仪表盘"""
        try:
            logger.info("生成交互式仪表盘")
            
            # 时间趋势分析
            time_df = df.resample('6H', on='CommentTime').agg({
                'CommentId': 'count',
                'sentiment': 'mean'
            }).reset_index()
            
            # 情感分布
            sentiment_dist = df['sentiment_label'].value_counts()
            
            fig = make_subplots(
                rows=2, cols=2,
                specs=[[{"type": "scatter"}, {"type": "pie"}],
                       [{"type": "treemap"}, {"type": "bar"}]],
                subplot_titles=(
                    "评论时间趋势", "情感分布", 
                    "地区热度", "活跃用户TOP10"
                )
            )
            
            # 时间趋势
            fig.add_trace(
                go.Scatter(
                    x=time_df['CommentTime'],
                    y=time_df['CommentId'],
                    mode='lines+markers',
                    name='评论量'
                ),
                row=1, col=1
            )
            
            # 情感分布
            fig.add_trace(
                go.Pie(
                    labels=sentiment_dist.index,
                    values=sentiment_dist.values,
                    hole=0.3
                ),
                row=1, col=2
            )
            
            # 地区分布
            region_df = df.groupby('CommentIp', as_index=False).agg(
                comment_count=('CommentId', 'count'),
                avg_sentiment=('sentiment', 'mean')
            )
            if not region_df.empty:
                # 计算占比
                region_df['percentage'] = region_df['comment_count'] / region_df['comment_count'].sum()
                fig.add_trace(
                    go.Treemap(
                        labels=region_df['CommentIp'],
                        parents=['']*len(region_df),
                        values=region_df['percentage'],  # 使用占比值
                        marker_colorscale='RdYlGn',
                        textinfo="label+percent entry"  # 只显示标签和百分比
                    ),
                    row=2, col=1
                )
            
            # 活跃用户
            user_df = df.groupby('CommentUserId', as_index=False)\
                      .agg(comment_count=('CommentId', 'count'))\
                      .nlargest(10, 'comment_count')
            fig.add_trace(
                go.Bar(
                    x=user_df['CommentUserId'],
                    y=user_df['comment_count'],
                    marker_color='skyblue'
                ),
                row=2, col=2
            )
            
            fig.update_layout(
                height=1200,
                width=1600,
                title_text="舆情分析仪表盘",
                showlegend=False
            )
            fig.write_html(self.config.output_dir / "dashboard.html")
        except Exception as e:
            logger.error(f"仪表盘生成失败: {str(e)}")

class ReportGenerator:
    """报告生成器"""
    def __init__(self):
        self.config = AnalysisConfig()
        self.doc = Document()
        
    def add_cover(self):
        """报告封面"""
        self.doc.add_heading('社交媒体舆情分析报告', 0)
        self.doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        self.doc.add_paragraph("系统版本：v2.1")
        
    def add_summary(self, df: pd.DataFrame):
        """数据概览"""
        self.doc.add_heading('数据概览', level=1)
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'LightShading'
        
        headers = table.rows[0].cells
        headers[0].text = '总评论数'
        headers[1].text = '时间范围'
        headers[2].text = '有效用户数'
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(len(df))
        row_cells[1].text = f"{df['CommentTime'].min():%Y-%m-%d} 至 {df['CommentTime'].max():%Y-%m-%d}"
        row_cells[2].text = str(df['CommentUserId'].nunique())
        
    def analyze_comments_by_region(self, df: pd.DataFrame):
        try:
            logger.info("按地区分析评论")
            # 地区评论占比数量
            region_count = df['CommentIp'].value_counts(normalize=True)
            region_count.to_csv(self.config.output_dir / 'region_comment_ratio.csv')
            
            # 地区评论词云图
            for region in df['CommentIp'].unique():
                region_df = df[df['CommentIp'] == region]
                text = ' '.join(region_df['clean_text'])
                if not text:
                    logger.warning(f"地区 {region} 无有效文本生成词云")
                    continue
                wc = WordCloud(
                    font_path=self.config.font_path,
                    width=1600,
                    height=1200,
                    background_color='white',
                    collocations=False
                ).generate(text)
                plt.figure(figsize=(20, 15))
                plt.imshow(wc)
                plt.axis('off')
                # 计算该地区评论数量
                plt.rcParams['font.family'] = 'SimHei' 
                # 添加标题
                title = f'{region}'
                plt.title(title, fontsize=28, color='red')

                plt.savefig(self.config.output_dir / f'wordcloud_{region}.png', bbox_inches='tight')
                plt.close()
            
            # 总体词云图
            self.generate_wordcloud(df)
            
            # 高频用户评论内容
            user_comment_count = df.groupby('CommentUserId')['CommentId'].count()
            high_frequency_users = user_comment_count.nlargest(10).index
            high_frequency_comments = df[df['CommentUserId'].isin(high_frequency_users)][['CommentUserId', 'clean_text']]
            high_frequency_comments.to_csv(self.config.output_dir / 'high_frequency_comments.csv')
        except Exception as e:
            logger.error(f"地区评论分析失败: {str(e)}")
    
    def generate(self, df: pd.DataFrame):
        """生成完整报告"""
        try:
            logger.info("生成分析报告")
            self.add_cover()
            self.add_summary(df)
            
            # 按地区分析评论
            self.analyze_comments_by_region(df)
            
            # 添加图表
            self._add_image_section('用户行为分析', 'time_distribution.png')
            self._add_image_section('情感分析', 'sentiment_dist.png')
            self._add_image_section('地区分布', 'region_dist.png')
            self._add_image_section('关键词云', 'wordcloud.png')
            
            # 添加地区评论词云图
            for region in df['CommentIp'].unique():
                self._add_image_section(f'地区 {region} 评论词云图', f'wordcloud_{region}.png')
            
            # 保存报告
            self.doc.save(self.config.output_dir / 'analysis_report.docx')
            logger.info("报告生成完成")
        except Exception as e:
            logger.error(f"报告生成失败: {str(e)}")
            
    def _add_image_section(self, title: str, image_path: str):
        """添加带图表的章节"""
        self.doc.add_heading(title, level=2)
        full_path = self.config.output_dir / image_path
        if full_path.exists():
            self.doc.add_picture(str(full_path), width=Inches(6))
        else:
            self.doc.add_paragraph(f"图表缺失：{image_path}")
        self.doc.add_page_break()

class PublicOpinionMonitor:
    """舆情监控主控制器"""
    def __init__(self, db_config: dict):
        self.config = AnalysisConfig()
        self.db_config = db_config
        self.df = None
        
    def execute_pipeline(self):
        """执行分析流水线"""
        try:
            logger.info("启动舆情分析流水线")
            
            # 数据准备
            processor = DataProcessor(self.db_config)
            raw_df = processor.preprocess()
            
            # 情感分析
            sentiment_analyzer = SentimentAnalyzer()
            analyzed_df = sentiment_analyzer.analyze(raw_df)
            
            # 可视化
            viz_engine = VisualizationEngine()
            viz_engine.generate_wordcloud(analyzed_df)
            viz_engine.create_dashboard(analyzed_df)
            
            # 报告生成
            report_gen = ReportGenerator()
            report_gen.generate(analyzed_df)
            
            # 数据持久化
            self._save_results(analyzed_df)
            
            logger.info("分析流水线成功完成")
            return True
        except Exception as e:
            logger.critical(f"分析流水线失败: {str(e)}")
            return False
            
    def _save_results(self, df: pd.DataFrame):
        """保存分析结果"""
        try:
            # 保存处理数据
            df.to_parquet(self.config.output_dir / 'processed_data.parquet')
            
            # 保存元数据
            meta = {
                'process_time': datetime.now().isoformat(),
                'data_version': 'v2.1',
                'record_count': len(df)
            }
            with open(self.config.output_dir / 'metadata.json', 'w') as f:
                json.dump(meta, f)
        except Exception as e:
            logger.error(f"结果保存失败: {str(e)}")

if __name__ == "__main__":
    # 数据库配置
    db_config = {
        'host': 'localhost',
        'port': 3306,
        'user': 'root',
        'password': 'Aa123Aa',
        'db': 'self_other_project'
    }
    
    # 执行分析
    monitor = PublicOpinionMonitor(db_config)
    success = monitor.execute_pipeline()
    
    if success:
        logger.info("舆情分析成功完成")
        sys.exit(0)
    else:
        logger.error("舆情分析执行失败")
        sys.exit(1)